"""
Markdown to Word converter for the DiDi CX deliverables.

Covers the subset of Markdown used in this project: headings, paragraphs, tables,
fenced code blocks, bullet and numbered lists, checklists, blockquotes and the
inline marks for bold, italic and code.

Straight quotes are preserved so DAX pasted out of the Word file still compiles.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── DiDi palette ──────────────────────────────────────────────────────────────
ORANGE = RGBColor(0xFF, 0x66, 0x00)
NAVY = RGBColor(0x0B, 0x1F, 0x33)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x64, 0x74, 0x8B)
CODE_FG = RGBColor(0x1F, 0x2A, 0x37)

NAVY_HEX = "0B1F33"
ORANGE_HEX = "FF6600"
CODE_BG_HEX = "F5F6F8"
BORDER_HEX = "D9DDE3"
ROW_ALT_HEX = "FAFBFC"

BODY_FONT = "Segoe UI"
CODE_FONT = "Consolas"

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")


# ── low level helpers ─────────────────────────────────────────────────────────

def _shade(element, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    element.append(shd)


def shade_cell(cell, hex_color: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), hex_color)


def shade_paragraph(paragraph, hex_color: str) -> None:
    _shade(paragraph._p.get_or_add_pPr(), hex_color)


def set_paragraph_border(paragraph, *, left: str | None = None, bottom: str | None = None,
                         size: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge, color in (("left", left), ("bottom", bottom)):
        if not color:
            continue
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), color)
        borders.append(el)
    p_pr.append(borders)


def set_table_borders(table, hex_color: str = BORDER_HEX) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tbl_pr.append(borders)


def repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def add_page_number_footer(section, label: str) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{label}   |   ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.font.name = BODY_FONT

    run = footer.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.font.name = BODY_FONT
    for instr, kind in (("begin", "fldChar"), ("PAGE", "instrText"), ("end", "fldChar")):
        el = OxmlElement(f"w:{kind}")
        if kind == "fldChar":
            el.set(qn("w:fldCharType"), instr)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        run._r.append(el)


# ── inline formatting ─────────────────────────────────────────────────────────

def add_inline(paragraph, text: str, *, size: Pt = Pt(10), color: RGBColor = DARK,
               bold: bool = False, italic: bool = False, font: str = BODY_FONT) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        run = paragraph.add_run()
        run.font.size = size
        run.font.name = font
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic

        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = CODE_FONT
            run.font.size = Pt(size.pt - 0.5)
            run.font.color.rgb = CODE_FG
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        else:
            run.text = part


# ── block builders ────────────────────────────────────────────────────────────

def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format

    if level == 1:
        fmt.space_before = Pt(22)
        fmt.space_after = Pt(8)
        add_inline(p, text, size=Pt(18), color=NAVY, bold=True)
        set_paragraph_border(p, bottom=ORANGE_HEX, size=12)
    elif level == 2:
        fmt.space_before = Pt(18)
        fmt.space_after = Pt(6)
        add_inline(p, text, size=Pt(13.5), color=NAVY, bold=True)
    elif level == 3:
        fmt.space_before = Pt(13)
        fmt.space_after = Pt(4)
        add_inline(p, text, size=Pt(11.5), color=DARK, bold=True)
    else:
        fmt.space_before = Pt(10)
        fmt.space_after = Pt(3)
        add_inline(p, text, size=Pt(10.5), color=MUTED, bold=True)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    add_inline(p, text)


def add_bullet(doc: Document, text: str, ordered: bool, indent_level: int = 0) -> None:
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.75 + 0.6 * indent_level)
    add_inline(p, text)


def add_checkbox(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run("\u2610  ")
    run.font.size = Pt(11)
    run.font.name = BODY_FONT
    add_inline(p, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for index, line in enumerate(lines):
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.space_before = Pt(6 if index == 0 else 0)
        fmt.space_after = Pt(6 if index == len(lines) - 1 else 0)
        fmt.left_indent = Cm(0.4)
        fmt.line_spacing = 1.0
        shade_paragraph(p, CODE_BG_HEX)
        set_paragraph_border(p, left=ORANGE_HEX, size=12)
        run = p.add_run(line if line else " ")
        run.font.name = CODE_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = CODE_FG


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(8)
    fmt.left_indent = Cm(0.4)
    fmt.line_spacing = 1.2
    shade_paragraph(p, "FFF7ED")
    set_paragraph_border(p, left=ORANGE_HEX, size=18)
    add_inline(p, text, size=Pt(9.5), color=DARK)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    header, *body = rows
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)

    # Column widths proportional to the longest content in each column
    weights = []
    for col in range(len(header)):
        longest = max(len(r[col]) for r in rows if col < len(r))
        weights.append(max(longest, 6))
    total = sum(weights)
    available = 16.2  # cm inside the margins

    for cell, text, weight in zip(table.rows[0].cells, header, weights):
        cell.width = Cm(available * weight / total)
        shade_cell(cell, NAVY_HEX)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_inline(p, text, size=Pt(8.5), color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    repeat_header_row(table.rows[0])

    for index, row_data in enumerate(body):
        cells = table.add_row().cells
        for col, cell in enumerate(cells):
            cell.width = Cm(available * weights[col] / total)
            if index % 2 == 1:
                shade_cell(cell, ROW_ALT_HEX)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            value = row_data[col] if col < len(row_data) else ""
            add_inline(p, value, size=Pt(8.5))


# ── parser ────────────────────────────────────────────────────────────────────

def _split_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|?[\s:\-|]+\|?", line.strip())) and "-" in line


def render_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        # fenced code
        if line.startswith("```"):
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i].rstrip())
                i += 1
            i += 1
            add_code_block(doc, block)
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line):
            i += 1
            continue

        # table
        if line.startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = _split_table_row(line)
            i += 2
            body = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body.append(_split_table_row(lines[i]))
                i += 1
            add_table(doc, [header] + body)
            continue

        # heading
        heading = re.match(r"(#{1,6})\s+(.*)", line)
        if heading:
            add_heading(doc, heading.group(2).strip(), len(heading.group(1)))
            i += 1
            continue

        # blockquote
        if line.startswith(">"):
            chunk = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                chunk.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_callout(doc, " ".join(chunk))
            continue

        # checklist
        checkbox = re.match(r"[-*]\s+\[[ xX]\]\s+(.*)", line)
        if checkbox:
            add_checkbox(doc, checkbox.group(1).strip())
            i += 1
            continue

        # bullets and numbered lists
        bullet = re.match(r"([-*])\s+(.*)", line)
        numbered = re.match(r"(\d+)\.\s+(.*)", line)
        if bullet or numbered:
            indent = (len(raw) - len(raw.lstrip())) // 3
            text = (bullet or numbered).group(2).strip()
            add_bullet(doc, text, ordered=bool(numbered), indent_level=indent)
            i += 1
            continue

        add_paragraph(doc, line)
        i += 1


# ── document shell ────────────────────────────────────────────────────────────

def build_document(title: str, subtitle: str, footer_label: str) -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = DARK

    brand = doc.add_paragraph()
    brand.paragraph_format.space_after = Pt(0)
    run = brand.add_run("DiDi")
    run.font.name = BODY_FONT
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = ORANGE

    tag = doc.add_paragraph()
    tag.paragraph_format.space_after = Pt(16)
    run = tag.add_run("CX SERVICE OPERATIONS  |  QUALITY ANALYTICS")
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.bold = True
    run.font.color.rgb = MUTED

    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(title)
    run.font.name = BODY_FONT
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    run = sub.add_run(subtitle)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED
    set_paragraph_border(sub, bottom=ORANGE_HEX, size=12)

    add_page_number_footer(section, footer_label)
    return doc


def convert(md_path: Path, out_path: Path, title: str, subtitle: str, footer_label: str) -> None:
    doc = build_document(title, subtitle, footer_label)
    markdown = md_path.read_text(encoding="utf-8")

    # The first H1 becomes the cover title, so drop it from the body
    markdown = re.sub(r"\A#\s+.*\n", "", markdown, count=1)

    render_markdown(doc, markdown)
    doc.save(out_path)
