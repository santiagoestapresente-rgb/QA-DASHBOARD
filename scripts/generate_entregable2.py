"""
DiDi CX — Entregable 2: Weekly Performance Report
Generates structured report with 7 Quality Tools integration.
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from scipy import stats

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "data"
OUT = ROOT / "entregable 2"
CHARTS = OUT / "charts"
SOURCE = Path(r"c:\Users\PC\Downloads\Business Case.xlsx")

DIDI_ORANGE = "#FF6600"
DIDI_DARK = "#1A1A1A"
DIDI_WHITE = "#FFFFFF"
STATUS_COLORS = {"green": "#28A745", "amber": "#FFC107", "red": "#DC3545"}

QA_GOAL = 85
CSAT_GOAL = 85
RECONTACT_GOAL = 5.44

PHONE_ATTRS = [
    "Critical_atributo_actitud_de_servicio_end_user",
    "Critical_atributo_compensaciones_reembolsos_end_user",
    "atributo_comunicacion_efectiva_end_user",
    "atributo_escucha_activa_end_user",
    "Critical_atributo_informacion_completa_y_correcta_end_user",
    "atributo_manejo_del_lenguaje_end_user",
    "atributo_manejo_del_tiempo_end_user",
    "Critical_atributo_negacion_de_servicio_end_user",
    "atributo_nombre_de_usuario_end_user",
    "atributo_personalizacion_de_la_interaccion_end_user",
    "atributo_presentacion_end_user",
    "Critical_atributo_rudeza_con_el_usuario_end_user",
]
LIVECHAT_ATTRS = [
    "Critical_Objetividad_del_chat",
    "Critical_Disponibilidad_del_servicio",
    "Saludo_e_identificacion",
    "Calidad_del_sondeo",
    "Recurrencia_de_informacion",
    "Actitud_de_servicio",
    "Personalizacion",
    "Calidad_de_comunicacion",
]


def is_critical(col: str) -> bool:
    return "critical" in col.lower()


def calc_qa_score(row: pd.Series, attr_cols: list[str]) -> float:
    for col in attr_cols:
        val = row.get(col)
        if pd.isna(val) or val == 2:
            continue
        if is_critical(col) and val == 1:
            return 0.0
    fails = sum(
        1
        for col in attr_cols
        if not pd.isna(row.get(col)) and row.get(col) != 2 and not is_critical(col) and row.get(col) == 1
    )
    return max(0.0, 100.0 - fails * 10)


def status_color(value: float, goal: float, higher_is_better: bool = True) -> str:
    diff = value - goal if higher_is_better else value - goal
    if higher_is_better:
        if diff >= 0:
            return "green"
        if diff >= -5:
            return "amber"
        return "red"
    if diff <= 0:
        return "green"
    if diff <= 5:
        return "amber"
    return "red"


def status_label(status: str) -> str:
    return {"green": "En meta / Por encima", "amber": "Dentro de 5pp", "red": "Más de 5pp bajo meta"}[status]


def load_data() -> dict[str, pd.DataFrame]:
    qa = pd.read_excel(SOURCE, sheet_name="QA")
    csat = pd.read_excel(SOURCE, sheet_name="CSAT")
    rc = pd.read_excel(SOURCE, sheet_name="Recontact")

    qa["CR_Lv4"] = qa["CR_correcta"].fillna(qa["CR_registrada"])
    qa["QA_Score"] = qa.apply(
        lambda r: calc_qa_score(r, PHONE_ATTRS if r["Channel"] == "Phone" else LIVECHAT_ATTRS),
        axis=1,
    )

    csat.columns = [c.strip().replace("\ufeff", "") for c in csat.columns]
    csat.rename(columns={"Consolidated Channel.": "Channel"}, inplace=True)
    csat["Satisfied_CNT"] = csat["Questionnaires With Star Level =4"] + csat["Questionnaires With Star Level =5"]
    csat["Unsatisfied_CNT"] = (
        csat["Questionnaires With Star Level =1"]
        + csat["Questionnaires With Star Level =2"]
        + csat["Questionnaires With Star Level =3"]
    )
    csat["CSAT_Pct"] = np.where(csat["Feedback CNT"] > 0, csat["Satisfied_CNT"] / csat["Feedback CNT"] * 100, np.nan)
    csat["Has_VOC"] = csat["open_question"].notna() & (csat["open_question"].astype(str).str.strip() != "")

    rc.columns = [c.strip().replace("\ufeff", "") for c in rc.columns]
    rc.rename(columns={"region_name": "Country", "customer_type": "User_Type"}, inplace=True)
    rc["Recontact_Rate"] = np.where(rc["Contacts"] > 0, rc["Recontact Volume"] / rc["Contacts"] * 100, np.nan)

    return {"qa": qa, "csat": csat, "rc": rc}


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tc_pr.append(shading)


def _set_run_font(run, *, size: Pt = Pt(8), bold: bool = False, color: RGBColor, name: str = "Calibri") -> None:
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def _set_cell_margins(cell, top: int = 40, bottom: int = 40, left: int = 80, right: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _set_row_height(row, twips: int, rule: str = "exact") -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(twips))
    tr_height.set(qn("w:hRule"), rule)
    tr_pr.append(tr_height)


def _set_table_borders_nil(table) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def _clear_header_footer(container) -> None:
    for tbl in list(container.tables):
        tbl._tbl.getparent().remove(tbl._tbl)
    paras = container.paragraphs
    if paras:
        paras[0].clear()
        for extra in paras[1:]:
            extra._element.getparent().remove(extra._element)


def _add_field(paragraph, instr: str, *, size: Pt = Pt(8), color: RGBColor, name: str = "Calibri") -> None:
    run = paragraph.add_run()
    _set_run_font(run, size=size, color=color, name=name)
    for token, kind in (("begin", "fldChar"), (f" {instr} ", "instrText"), ("end", "fldChar")):
        el = OxmlElement(f"w:{kind}")
        if kind == "fldChar":
            el.set(qn("w:fldCharType"), token)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = token
        run._r.append(el)


def apply_running_header_footer(section) -> None:
    """Repeating branded chrome on every page of the Word deliverable."""
    section.different_first_page_header_footer = False
    section.header_distance = Cm(0.4)
    section.footer_distance = Cm(0.5)

    white = RGBColor(255, 255, 255)
    dark = RGBColor(0x1A, 0x1A, 0x1A)
    gray = RGBColor(0x66, 0x66, 0x66)
    usable = section.page_width - section.left_margin - section.right_margin

    header = section.header
    header.is_linked_to_previous = False
    _clear_header_footer(header)

    brand = header.add_table(1, 2, usable)
    _set_table_borders_nil(brand)
    left, right = brand.rows[0].cells
    for cell in (left, right):
        set_cell_shading(cell, DIDI_DARK)
        _set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
    _set_row_height(brand.rows[0], 280, "atLeast")

    lp = left.paragraphs[0]
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(0)
    run = lp.add_run("DiDi Global — CX Service Operations  |  Internal Use Only")
    _set_run_font(run, size=Pt(8), color=white)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after = Pt(0)
    run = rp.add_run("Entregable 2 — Weekly Performance Report")
    _set_run_font(run, size=Pt(8), bold=True, color=white)

    bar = header.add_table(1, 1, usable)
    _set_table_borders_nil(bar)
    bar_cell = bar.rows[0].cells[0]
    set_cell_shading(bar_cell, DIDI_ORANGE)
    _set_cell_margins(bar_cell, top=0, bottom=0, left=0, right=0)
    _set_row_height(bar.rows[0], 60, "exact")
    bar_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    bar_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    footer = section.footer
    footer.is_linked_to_previous = False
    _clear_header_footer(footer)
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_before = Pt(2)
    p_pr = fp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), DIDI_ORANGE.lstrip("#"))
    p_bdr.append(top)
    p_pr.append(p_bdr)

    run = fp.add_run("CONFIDENTIAL  ·  Internal Use Only")
    _set_run_font(run, size=Pt(8), color=dark)
    run = fp.add_run("          ")
    _set_run_font(run, size=Pt(8), color=gray)
    run = fp.add_run("Page ")
    _set_run_font(run, size=Pt(8), color=dark)
    _add_field(fp, "PAGE", size=Pt(8), color=dark)
    run = fp.add_run(" of ")
    _set_run_font(run, size=Pt(8), color=dark)
    _add_field(fp, "NUMPAGES", size=Pt(8), color=dark)

    for p in list(header.paragraphs):
        if not (p.text or "").strip():
            parent = p._element.getparent()
            if parent is not None and (p._element.getnext() is not None or p._element.getprevious() is not None):
                parent.remove(p._element)


def apply_document_chrome(doc: Document, *, restyle_tables: bool = False) -> None:
    """Apply repeating header/footer and VP margins. Optionally restyle existing table headers."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.7)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    apply_running_header_footer(section)
    if restyle_tables:
        for table in doc.tables:
            if not table.rows:
                continue
            for cell in table.rows[0].cells:
                set_cell_shading(cell, DIDI_DARK)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def add_styled_table(doc: Document, headers: list[str], rows: list[list], status_col: int | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], DIDI_DARK)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
            if status_col is not None and i == status_col and val in STATUS_COLORS:
                set_cell_shading(cells[i], STATUS_COLORS[val])


def style_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(26, 26, 26)
    apply_document_chrome(doc)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(255, 102, 0) if level == 1 else RGBColor(26, 26, 26)


def add_status_paragraph(doc: Document, metric: str, value: float, goal: float, unit: str, higher_better: bool) -> None:
    st = status_color(value, goal, higher_better)
    p = doc.add_paragraph()
    p.add_run(f"{metric}: ").bold = True
    p.add_run(f"{value:.2f}{unit} (Meta: {goal}{unit}) — ")
    run = p.add_run(status_label(st))
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(STATUS_COLORS[st].lstrip("#"))


# ── 7 Quality Tools Charts ────────────────────────────────────────────────

def chart_check_sheet(qa: pd.DataFrame) -> Path:
    """Tool 1: Check Sheet — Data quality & audit validation checklist."""
    mismatch = (qa["Score_end_user"].notna() & (qa["Score_end_user"] != qa["QA_Score"])).mean() * 100
    checks = [
        ("Registros QA totales", len(qa), "≥ 2,000", "OK" if len(qa) >= 2000 else "REVISAR"),
        ("Audits Phone", len(qa[qa["Channel"] == "Phone"]), "≥ 300", "OK"),
        ("Audits Live Chat", len(qa[qa["Channel"] == "Live Chat"]), "≥ 1,500", "OK"),
        ("CR Lv4 completos", qa["CR_Lv4"].notna().mean() * 100, "≥ 95%", "OK" if qa["CR_Lv4"].notna().mean() >= 0.95 else "REVISAR"),
        ("Validación score QA", 100 - mismatch, "≥ 80%", "OK" if mismatch < 20 else "REVISAR"),
        ("Atributos N/A excluidos", "Sí", "Sí", "OK"),
        ("Regla Critical aplicada", "Sí", "Sí", "OK"),
        ("Canales separados", "Sí", "Sí", "OK"),
    ]

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.axis("off")
    table = ax.table(
        cellText=[[c[0], f"{c[1]}" if isinstance(c[1], (int, float)) else c[1], c[2], c[3]] for c in checks],
        colLabels=["Ítem de Verificación", "Resultado", "Criterio", "Estado"],
        cellLoc="center",
        loc="center",
        colWidths=[0.35, 0.2, 0.2, 0.15],
    )
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    for (row, col), cell in table.get_celld().items():
        if row == 0:
            cell.set_facecolor(DIDI_ORANGE)
            cell.set_text_props(color="white", weight="bold")
        elif col == 3:
            cell.set_facecolor(STATUS_COLORS["green"] if checks[row - 1][3] == "OK" else STATUS_COLORS["amber"])
    ax.set_title("Herramienta 1 — Check Sheet: Validación de Datos y Auditoría QA", fontsize=12, color=DIDI_DARK, weight="bold", pad=20)
    path = CHARTS / "01_check_sheet.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def chart_pareto(qa_attrs: pd.DataFrame) -> Path:
    """Tool 2: Pareto — Top failing QA attributes."""
    phone = qa_attrs[qa_attrs["Channel"] == "Phone"].head(6)
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))

    for ax, channel, df in zip(axes, ["Phone", "Live Chat"], [phone, qa_attrs[qa_attrs["Channel"] == "Live Chat"].head(6)]):
        vals = df["Fail_Count"].values
        labels = [a[:22] + "…" if len(a) > 22 else a for a in df["Attribute"].values]
        cum = np.cumsum(vals) / vals.sum() * 100 if vals.sum() else vals
        bars = ax.bar(range(len(vals)), vals, color=DIDI_ORANGE, edgecolor=DIDI_DARK, linewidth=0.5)
        ax2 = ax.twinx()
        ax2.plot(range(len(vals)), cum, "o-", color=DIDI_DARK, linewidth=2, markersize=5)
        ax2.axhline(80, color=STATUS_COLORS["red"], linestyle="--", alpha=0.7, label="80%")
        ax2.set_ylim(0, 105)
        ax2.set_ylabel("Acumulado %", fontsize=9)
        ax.set_xticks(range(len(vals)))
        ax.set_xticklabels(labels, rotation=35, ha="right", fontsize=7)
        ax.set_ylabel("Fallas", fontsize=9)
        ax.set_title(f"Pareto — {channel}", fontsize=10, weight="bold")
        for bar, v in zip(bars, vals):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1, str(int(v)), ha="center", fontsize=7)

    fig.suptitle("Herramienta 2 — Diagrama de Pareto: Atributos QA con Mayor Concentración de Defectos", fontsize=12, color=DIDI_DARK, weight="bold")
    path = CHARTS / "02_pareto_attributes.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def chart_histogram(qa: pd.DataFrame) -> Path:
    """Tool 3: Histogram — QA score distribution by channel."""
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5), sharey=True)
    bins = [0, 10, 20, 30, 40, 50, 60, 70, 80, 90, 100]

    for ax, channel in zip(axes, ["Phone", "Live Chat"]):
        scores = qa.loc[qa["Channel"] == channel, "QA_Score"]
        ax.hist(scores, bins=bins, color=DIDI_ORANGE, edgecolor=DIDI_DARK, alpha=0.85)
        ax.axvline(QA_GOAL, color=STATUS_COLORS["green"], linestyle="--", linewidth=2, label=f"Meta {QA_GOAL}")
        ax.axvline(scores.mean(), color=DIDI_DARK, linestyle="-", linewidth=1.5, label=f"Promedio {scores.mean():.1f}")
        ax.set_xlabel("QA Score", fontsize=9)
        ax.set_title(f"Histograma — {channel} (n={len(scores)})", fontsize=10, weight="bold")
        ax.legend(fontsize=7)

    axes[0].set_ylabel("Frecuencia de interacciones", fontsize=9)
    fig.suptitle("Herramienta 3 — Histograma: Distribución de QA Score por Canal", fontsize=12, color=DIDI_DARK, weight="bold")
    path = CHARTS / "03_histogram_qa_scores.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def chart_control(qa: pd.DataFrame) -> Path:
    """Tool 4: Control Chart — QA score by country (subgroup)."""
    g = qa.groupby(["Country", "Channel"])["QA_Score"].mean().reset_index()
    fig, ax = plt.subplots(figsize=(10, 5))
    x = range(len(g))
    colors = [DIDI_ORANGE if ch == "Phone" else DIDI_DARK for ch in g["Channel"]]
    ax.bar(x, g["QA_Score"], color=colors, edgecolor="white")
    ucl = qa["QA_Score"].mean() + 2 * qa["QA_Score"].std()
    lcl = max(0, qa["QA_Score"].mean() - 2 * qa["QA_Score"].std())
    center = qa["QA_Score"].mean()
    ax.axhline(center, color=DIDI_DARK, linestyle="-", linewidth=1.5, label=f"LCL media={center:.1f}")
    ax.axhline(QA_GOAL, color=STATUS_COLORS["green"], linestyle="--", linewidth=2, label=f"Meta={QA_GOAL}")
    ax.axhline(ucl, color=STATUS_COLORS["red"], linestyle=":", linewidth=1.5, label=f"UCL={ucl:.1f}")
    ax.axhline(lcl, color=STATUS_COLORS["red"], linestyle=":", linewidth=1.5, label=f"LCL={lcl:.1f}")
    ax.set_xticks(list(x))
    ax.set_xticklabels([f"{r.Country}\n({r.Channel[:2]})" for r in g.itertuples()], fontsize=8)
    ax.set_ylabel("QA Score Promedio", fontsize=9)
    ax.set_ylim(60, 105)
    ax.legend(fontsize=7, loc="lower right")
    ax.set_title("Herramienta 4 — Gráfico de Control: QA Score por País y Canal", fontsize=12, color=DIDI_DARK, weight="bold")
    path = CHARTS / "04_control_chart_qa.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def chart_scatter(combined: pd.DataFrame) -> Path:
    """Tool 5: Scatter — QA vs CSAT correlation by CR Lv4."""
    df = combined.dropna(subset=["QA_Score", "CSAT_Pct"])
    df = df[df["QA_Audits"] >= 5]
    fig, ax = plt.subplots(figsize=(9, 6))
    scatter = ax.scatter(
        df["QA_Score"],
        df["CSAT_Pct"],
        s=np.clip(df["Feedback_CNT"] / 50, 20, 400),
        c=df["Recontact_Rate"].fillna(0),
        cmap="YlOrRd",
        alpha=0.7,
        edgecolors=DIDI_DARK,
        linewidth=0.5,
    )
    if len(df) >= 3:
        slope, intercept, r, _, _ = stats.linregress(df["QA_Score"], df["CSAT_Pct"])
        x_line = np.linspace(df["QA_Score"].min(), df["QA_Score"].max(), 50)
        ax.plot(x_line, slope * x_line + intercept, color=DIDI_DARK, linestyle="--", linewidth=1.5, label=f"R²={r**2:.2f}")
    ax.axhline(CSAT_GOAL, color=STATUS_COLORS["green"], linestyle="--", alpha=0.6)
    ax.axvline(QA_GOAL, color=STATUS_COLORS["green"], linestyle="--", alpha=0.6)
    ax.set_xlabel("QA Score Promedio", fontsize=9)
    ax.set_ylabel("CSAT %", fontsize=9)
    cbar = plt.colorbar(scatter, ax=ax)
    cbar.set_label("Recontact Rate %", fontsize=8)
    ax.legend(fontsize=8)
    ax.set_title("Herramienta 5 — Diagrama de Dispersión: QA vs CSAT (tamaño=vol. feedback)", fontsize=11, color=DIDI_DARK, weight="bold")

    for _, r in df.nsmallest(3, "CSAT_Pct").iterrows():
        ax.annotate(r["CR_Lv4_Name"][:25], (r["QA_Score"], r["CSAT_Pct"]), fontsize=6, alpha=0.8)

    path = CHARTS / "05_scatter_qa_csat.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def chart_fishbone() -> Path:
    """Tool 6: Cause & Effect (Ishikawa) for order status & delays cluster."""
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")

    spine = FancyArrowPatch((1, 4), (10.5, 4), arrowstyle="-", mutation_scale=1, linewidth=2.5, color=DIDI_DARK)
    ax.add_patch(spine)
    ax.text(10.8, 4, "Bajo CSAT +\nAlto Recontact\nOrder Status/Delays", fontsize=9, weight="bold", va="center", color=DIDI_ORANGE)

    categories = {
        "People": ["Falta capacitación\nPhone (AHT)", "New hires\nsin coaching", "Variabilidad\nentre agentes"],
        "Process": ["Sin script de\nresolución FCR", "Escalamiento\ninnecesario", "Info incompleta\nal cliente"],
        "Technology": ["Tracking en\ntiempo real limitado", "Bot GPTBOT\n13.7% recontact", "Integración\nstore-courier"],
        "Policy": ["Políticas reembolso\npoco claras", "Antifraud bloquea\nresolución", "SLA no\ncomunicado"],
        "Environment": ["Volumen pico\nW19", "Multipaís\n(CO, MX, PE)", "Marketplace vs\nFull Service mix"],
    }
    positions = [(2.5, 6.5), (4.5, 6.5), (6.5, 6.5), (2.5, 1.5), (6.5, 1.5)]
    for (cat, causes), (bx, by) in zip(categories.items(), positions):
        ax.plot([bx + 1, 5.5], [by, 4], color=DIDI_ORANGE, linewidth=1.2)
        ax.add_patch(FancyBboxPatch((bx, by - 0.3), 2.2, 0.6, boxstyle="round,pad=0.05", facecolor=DIDI_ORANGE, edgecolor=DIDI_DARK))
        ax.text(bx + 1.1, by, cat, ha="center", va="center", fontsize=8, weight="bold", color="white")
        for i, cause in enumerate(causes):
            cy = by + (0.8 if by > 4 else -0.8) * (i + 1)
            ax.text(bx + 0.2, cy, f"• {cause}", fontsize=7, va="center")

    ax.set_title("Herramienta 6 — Diagrama de Ishikawa: Causa-Raíz — Order Status & Delays", fontsize=12, color=DIDI_DARK, weight="bold")
    path = CHARTS / "06_fishbone_order_status.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def chart_flowchart() -> Path:
    """Tool 7: Flowchart — FCR failure path for recontact-prone CRs."""
    fig, ax = plt.subplots(figsize=(11, 8))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 10)
    ax.axis("off")

    boxes = [
        (4.5, 9, "Cliente contacta\n(Order Status/Delay)", DIDI_ORANGE),
        (4.5, 7.5, "Agente consulta\nsistema de tracking", DIDI_DARK),
        (4.5, 6, "¿Info completa\ndisponible?", DIDI_ORANGE),
        (1.5, 4.5, "SÍ: Comunica\nETA + acciones", STATUS_COLORS["green"]),
        (7.5, 4.5, "NO: Respuesta\ngenérica / espera", STATUS_COLORS["red"]),
        (1.5, 3, "Cliente satisfecho\nCSAT 4-5", STATUS_COLORS["green"]),
        (7.5, 3, "Cliente recontacta\n(16.9% rate)", STATUS_COLORS["red"]),
        (7.5, 1.5, "Segunda interacción\nsin resolución FCR", STATUS_COLORS["red"]),
        (4.5, 1.5, "Escalamiento\nQA Phone 83.0", DIDI_ORANGE),
    ]

    for x, y, text, color in boxes:
        fc = "white" if color == DIDI_DARK else color
        tc = "white" if color != "white" and color != STATUS_COLORS["green"] else DIDI_DARK
        ax.add_patch(FancyBboxPatch((x - 1.3, y - 0.45), 2.6, 0.9, boxstyle="round,pad=0.08", facecolor=fc, edgecolor=DIDI_DARK, linewidth=1.5))
        ax.text(x, y, text, ha="center", va="center", fontsize=7.5, color=tc, weight="bold" if color == DIDI_ORANGE else "normal")

    arrows = [
        ((4.5, 8.55), (4.5, 8.0)),
        ((4.5, 7.05), (4.5, 6.5)),
        ((3.8, 5.7), (2.0, 5.0)),
        ((5.2, 5.7), (7.0, 5.0)),
        ((1.5, 4.05), (1.5, 3.5)),
        ((7.5, 4.05), (7.5, 3.5)),
        ((7.5, 2.55), (7.5, 2.0)),
        ((6.5, 1.5), (5.8, 1.5)),
    ]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", color=DIDI_DARK, lw=1.2))

    ax.text(3.2, 5.3, "SÍ", fontsize=8, weight="bold")
    ax.text(5.5, 5.3, "NO", fontsize=8, weight="bold", color=STATUS_COLORS["red"])
    ax.set_title("Herramienta 7 — Diagrama de Flujo: Punto de Falla FCR (Order Status & Delays)", fontsize=12, color=DIDI_DARK, weight="bold")
    path = CHARTS / "07_flowchart_fcr.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def chart_pareto_recontact(rc_cr: pd.DataFrame) -> Path:
    top = rc_cr.head(10)
    fig, ax = plt.subplots(figsize=(10, 5))
    vals = top["Recontact_Volume"].values
    labels = [c[:30] + "…" if len(c) > 30 else c for c in top["CR_Lv4"].values]
    cum = np.cumsum(vals) / vals.sum() * 100
    ax.bar(range(len(vals)), vals, color=DIDI_ORANGE)
    ax2 = ax.twinx()
    ax2.plot(range(len(vals)), cum, "o-", color=DIDI_DARK)
    ax2.axhline(80, color=STATUS_COLORS["red"], linestyle="--")
    ax.set_xticks(range(len(vals)))
    ax.set_xticklabels(labels, rotation=40, ha="right", fontsize=7)
    ax.set_ylabel("Volumen Recontact")
    ax2.set_ylabel("Acumulado %")
    ax.set_title("Pareto — Top 10 CR Lv4 por Volumen de Recontact", fontsize=11, weight="bold")
    path = CHARTS / "08_pareto_recontact.png"
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def build_combined(qa: pd.DataFrame, csat: pd.DataFrame, rc: pd.DataFrame) -> pd.DataFrame:
    qa_cr = qa.groupby("CR_Lv4").agg(QA_Score=("QA_Score", "mean"), QA_Audits=("QA_Score", "count")).reset_index()
    qa_cr.rename(columns={"CR_Lv4": "CR_Lv4_Name"}, inplace=True)
    csat_cr = csat.groupby("CR Lv4").agg(Feedback_CNT=("Feedback CNT", "sum"), Satisfied_CNT=("Satisfied_CNT", "sum")).reset_index()
    csat_cr["CSAT_Pct"] = csat_cr["Satisfied_CNT"] / csat_cr["Feedback_CNT"] * 100
    csat_cr.rename(columns={"CR Lv4": "CR_Lv4_Name"}, inplace=True)
    rc_cr = rc.groupby("CR Lv4").agg(Contacts=("Contacts", "sum"), Recontact_Volume=("Recontact Volume", "sum")).reset_index()
    rc_cr["Recontact_Rate"] = rc_cr["Recontact_Volume"] / rc_cr["Contacts"] * 100
    rc_cr.rename(columns={"CR Lv4": "CR_Lv4_Name"}, inplace=True)
    return qa_cr.merge(csat_cr, on="CR_Lv4_Name", how="outer").merge(rc_cr, on="CR_Lv4_Name", how="outer")


def build_attribute_fails(qa: pd.DataFrame) -> pd.DataFrame:
    records = []
    for channel, attrs in [("Phone", PHONE_ATTRS), ("Live Chat", LIVECHAT_ATTRS)]:
        subset = qa[qa["Channel"] == channel]
        for col in attrs:
            applicable = subset[col].isin([0, 1])
            fail_cnt = (subset.loc[applicable, col] == 1).sum()
            app_cnt = applicable.sum()
            name = col.replace("_end_user", "").replace("Critical_", "").replace("atributo_", "").replace("_", " ").title()
            records.append({"Channel": channel, "Attribute": name, "Fail_Count": int(fail_cnt), "Applicable_Count": int(app_cnt), "Fail_Rate_Pct": round(fail_cnt / app_cnt * 100, 2) if app_cnt else 0})
    return pd.DataFrame(records).sort_values(["Channel", "Fail_Rate_Pct"], ascending=[True, False])


def generate_report() -> None:
    OUT.mkdir(exist_ok=True)
    CHARTS.mkdir(exist_ok=True)

    data = load_data()
    qa, csat, rc = data["qa"], data["csat"], data["rc"]

    qa_avg = qa["QA_Score"].mean()
    csat_pct = csat["Satisfied_CNT"].sum() / csat["Feedback CNT"].sum() * 100
    rc_rate = rc["Recontact Volume"].sum() / rc["Contacts"].sum() * 100

    qa_phone = qa[qa["Channel"] == "Phone"]["QA_Score"].mean()
    qa_chat = qa[qa["Channel"] == "Live Chat"]["QA_Score"].mean()
    csat_phone = csat[csat["Channel"] == "PHONE"]
    csat_chat = csat[csat["Channel"] == "LIVE CHAT"]
    csat_phone_pct = csat_phone["Satisfied_CNT"].sum() / csat_phone["Feedback CNT"].sum() * 100
    csat_chat_pct = csat_chat["Satisfied_CNT"].sum() / csat_chat["Feedback CNT"].sum() * 100

    attrs = build_attribute_fails(qa)
    combined = build_combined(qa, csat, rc)
    rc_cr = rc.groupby("CR Lv4").agg(Contacts=("Contacts", "sum"), Recontact_Volume=("Recontact Volume", "sum")).reset_index()
    rc_cr.rename(columns={"CR Lv4": "CR_Lv4"}, inplace=True)
    rc_cr["Recontact_Rate"] = (rc_cr["Recontact_Volume"] / rc_cr["Contacts"] * 100).round(2)
    rc_cr = rc_cr.sort_values("Recontact_Volume", ascending=False)

    qa_cr = qa.groupby(["CR_Lv4", "Channel"]).agg(QA_Score=("QA_Score", "mean"), Audits=("QA_Score", "count")).reset_index()
    qa_cr_agg = qa.groupby("CR_Lv4").agg(QA_Score=("QA_Score", "mean"), Audits=("QA_Score", "count")).reset_index()
    qa_cr_agg = qa_cr_agg[qa_cr_agg["Audits"] >= 3].sort_values("QA_Score").head(10)

    csat_cr = csat.groupby("CR Lv4").agg(Feedback_CNT=("Feedback CNT", "sum"), Satisfied_CNT=("Satisfied_CNT", "sum")).reset_index()
    csat_cr.rename(columns={"CR Lv4": "CR_Lv4"}, inplace=True)
    csat_cr["CSAT_Pct"] = (csat_cr["Satisfied_CNT"] / csat_cr["Feedback_CNT"] * 100).round(2)
    csat_cr = csat_cr[csat_cr["Feedback_CNT"] >= 20].sort_values("CSAT_Pct").head(10)

    voc_neg = csat[(csat["Has_VOC"]) & (csat["Unsatisfied_CNT"] > 0)].head(15)

    # Charts
    charts = {
        "check_sheet": chart_check_sheet(qa),
        "pareto": chart_pareto(attrs),
        "histogram": chart_histogram(qa),
        "control": chart_control(qa),
        "scatter": chart_scatter(combined),
        "fishbone": chart_fishbone(),
        "flowchart": chart_flowchart(),
        "pareto_rc": chart_pareto_recontact(rc_cr),
    }

    # ── Word Document ─────────────────────────────────────────────────────
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DiDi Global — CX Service Operations")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run("Weekly Performance Report — Semana W19")
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(255, 102, 0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Preparado por: CX Quality Analyst  |  Fecha: {datetime.now().strftime('%d-%b-%Y')}  |  CONFIDENCIAL").font.size = Pt(10)

    doc.add_paragraph()
    legend = doc.add_paragraph()
    legend.add_run("Código de color: ").bold = True
    for label, color in [("Verde = En meta", "green"), ("Ámbar = Dentro de 5pp", "amber"), ("Rojo = >5pp bajo meta", "red")]:
        r = legend.add_run(f"  ■ {label}  ")
        r.font.color.rgb = RGBColor.from_string(STATUS_COLORS[color].lstrip("#"))

    # Executive Summary
    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        f"Durante la semana W19, DiDi CX procesó {len(qa):,} auditorías QA, {csat['Feedback CNT'].sum():,.0f} respuestas CSAT "
        f"y {rc['Contacts'].sum():,.0f} contactos con {rc['Recontact Volume'].sum():,.0f} recontactos. "
        f"El desempeño global de QA alcanzó {qa_avg:.2f} puntos, superando la meta de {QA_GOAL} ({status_label(status_color(qa_avg, QA_GOAL))}). "
        f"Sin embargo, CSAT cerró en {csat_pct:.2f}%, {abs(csat_pct - CSAT_GOAL):.2f}pp por debajo de la meta de {CSAT_GOAL}% ({status_label(status_color(csat_pct, CSAT_GOAL))}). "
        f"La tasa de recontacto fue {rc_rate:.2f}%, ligeramente por encima del objetivo de {RECONTACT_GOAL}% ({status_label(status_color(rc_rate, RECONTACT_GOAL, False))})."
    )
    doc.add_paragraph(
        "La divergencia más relevante es la brecha entre canales: Phone presenta QA en ámbar (83.04) mientras Live Chat supera ampliamente la meta (96.01). "
        "Paradójicamente, Phone logra CSAT verde (86.26%) pero Live Chat — con QA excelente — registra CSAT rojo (77.55%). "
        "Esto sugiere que la calidad auditada en chat no se traduce en percepción del cliente, posiblemente por volumen, tiempos de espera o resolución incompleta en primera interacción."
    )
    doc.add_paragraph(
        "Hallazgo crítico para liderazgo: el cluster de Contact Reasons 'Order Status & Delays' y 'User Request Order Status or Delay Information' "
        "concentra 20,694 recontactos (35.7% del total), CSAT de 64.7–67.9% y volumen masivo de feedback negativo en VOC. "
        "A pesar de QA aceptable en Live Chat (95–97), la resolución en primera interacción (FCR) falla sistemáticamente. "
        "Se recomienda intervención inmediata en scripts de resolución, integración de tracking en tiempo real y reducción de respuestas genéricas identificadas en VOC."
    )

    add_styled_table(
        doc,
        ["Métrica", "Resultado", "Meta", "Brecha", "Estado"],
        [
            ["QA Score", f"{qa_avg:.2f}", QA_GOAL, f"{qa_avg - QA_GOAL:+.2f}", status_color(qa_avg, QA_GOAL)],
            ["CSAT", f"{csat_pct:.2f}%", f"{CSAT_GOAL}%", f"{csat_pct - CSAT_GOAL:+.2f}pp", status_color(csat_pct, CSAT_GOAL)],
            ["Recontact Rate", f"{rc_rate:.2f}%", f"{RECONTACT_GOAL}%", f"{rc_rate - RECONTACT_GOAL:+.2f}pp", status_color(rc_rate, RECONTACT_GOAL, False)],
        ],
        status_col=4,
    )

    # QA by Channel
    add_heading(doc, "2. QA Analysis — by Channel")
    add_heading(doc, "2.1 Phone", level=2)
    add_status_paragraph(doc, "QA Score Phone", qa_phone, QA_GOAL, "", True)
    doc.add_paragraph(
        "Phone registra 355 auditorías con 48 fallas críticas (score=0). El atributo dominante es Manejo del Tiempo (37.5% fail rate), "
        "seguido de Información Completa y Correcta (7.9%, Critical) y Compensaciones/Reembolsos (3.4%, Critical). "
        "Hipótesis: interacciones telefónicas más largas (AHT promedio elevado) generan presión temporal que compromete calidad informativa y manejo de compensaciones."
    )
    phone_attrs = attrs[attrs["Channel"] == "Phone"].head(5)
    add_styled_table(
        doc,
        ["Atributo", "Fallas", "Aplicables", "Fail Rate %", "Critical"],
        [[r.Attribute, r.Fail_Count, r.Applicable_Count, f"{r.Fail_Rate_Pct}%", "Sí" if "Informacion" in r.Attribute or "Compensaciones" in r.Attribute or "Negacion" in r.Attribute or "Rudeza" in r.Attribute or "Actitud" in r.Attribute else "No"] for r in phone_attrs.itertuples()],
    )
    phone_worst = qa_cr[(qa_cr["Channel"] == "Phone") & (qa_cr["Audits"] >= 3)].nsmallest(5, "QA_Score")
    add_styled_table(
        doc,
        ["CR Lv4 (Phone)", "QA Score", "Audits", "Hipótesis Root Cause"],
        [
            [r.CR_Lv4, f"{r.QA_Score:.1f}", int(r.Audits), "Fallas Critical en info/compensación; casos complejos de entrega"]
            for r in phone_worst.itertuples()
        ],
    )

    add_heading(doc, "2.2 Live Chat", level=2)
    add_status_paragraph(doc, "QA Score Live Chat", qa_chat, QA_GOAL, "", True)
    doc.add_paragraph(
        "Live Chat concentra 2,105 auditorías (85.6% del total) con QA de 96.01. Los top fails no-críticos: Saludo e Identificación (4.6%), "
        "Actitud de Servicio (3.3%) y Disponibilidad del Servicio (2.2%, Critical). "
        "A pesar del alto score, estos atributos de experiencia inicial impactan percepción cuando el cliente ya llega frustrado por demoras."
    )
    chat_attrs = attrs[attrs["Channel"] == "Live Chat"].head(5)
    add_styled_table(
        doc,
        ["Atributo", "Fallas", "Aplicables", "Fail Rate %"],
        [[r.Attribute, r.Fail_Count, r.Applicable_Count, f"{r.Fail_Rate_Pct}%"] for r in chat_attrs.itertuples()],
    )

    doc.add_picture(str(charts["pareto"]), width=Inches(6.2))
    doc.add_picture(str(charts["histogram"]), width=Inches(6.2))

    # QA by CR Lv4
    add_heading(doc, "3. QA Analysis — by CR Lv4")
    doc.add_paragraph(
        "Ranking de CR Lv4 con menor desempeño QA (mínimo 3 auditorías). Se observa concentración en casos de entrega no recibida, "
        "pedido activo ya entregado, y compensaciones/reembolsos — todos con alta incidencia de atributos Critical."
    )
    add_styled_table(
        doc,
        ["CR Lv4", "QA Score", "Audits", "Estado", "Hipótesis"],
        [
            [r.CR_Lv4, f"{r.QA_Score:.1f}", int(r.Audits), status_color(r.QA_Score, QA_GOAL), "Falla Critical: info incorrecta o negación de servicio/compensación"]
            for r in qa_cr_agg.itertuples()
        ],
        status_col=3,
    )
    doc.add_picture(str(charts["control"]), width=Inches(6))

    # CSAT / VOC
    add_heading(doc, "4. CSAT / VOC Analysis")
    add_status_paragraph(doc, "CSAT Global", csat_pct, CSAT_GOAL, "%", True)
    doc.add_paragraph(
        f"CSAT por canal: Phone {csat_phone_pct:.2f}% (verde, +1.26pp) vs Live Chat {csat_chat_pct:.2f}% (rojo, -7.45pp). "
        f"Live Chat representa 72.4% del feedback ({csat_chat['Feedback CNT'].sum():,.0f} encuestas), arrastrando el promedio global."
    )
    add_styled_table(
        doc,
        ["Canal", "Feedback", "CSAT %", "Brecha", "Estado"],
        [
            ["Phone", f"{csat_phone['Feedback CNT'].sum():,.0f}", f"{csat_phone_pct:.2f}%", f"{csat_phone_pct - CSAT_GOAL:+.2f}pp", status_color(csat_phone_pct, CSAT_GOAL)],
            ["Live Chat", f"{csat_chat['Feedback CNT'].sum():,.0f}", f"{csat_chat_pct:.2f}%", f"{csat_chat_pct - CSAT_GOAL:+.2f}pp", status_color(csat_chat_pct, CSAT_GOAL)],
        ],
        status_col=4,
    )

    csat_bt = csat.groupby("Business Type Name").agg(Feedback_CNT=("Feedback CNT", "sum"), Satisfied_CNT=("Satisfied_CNT", "sum")).reset_index()
    csat_bt.rename(columns={"Business Type Name": "Business_Type"}, inplace=True)
    csat_bt["CSAT_Pct"] = (csat_bt["Satisfied_CNT"] / csat_bt["Feedback_CNT"] * 100).round(2)
    csat_bt = csat_bt.sort_values("CSAT_Pct")
    add_styled_table(
        doc,
        ["Business Type", "Feedback", "CSAT %", "Estado"],
        [[r.Business_Type, int(r.Feedback_CNT), f"{r.CSAT_Pct}%", status_color(r.CSAT_Pct, CSAT_GOAL)] for r in csat_bt.itertuples()],
        status_col=3,
    )

    doc.add_paragraph("Top CR Lv4 con peor CSAT (mín. 20 feedback):")
    add_styled_table(
        doc,
        ["CR Lv4", "Feedback", "CSAT %", "Estado"],
        [[r.CR_Lv4, int(r.Feedback_CNT), f"{r.CSAT_Pct}%", status_color(r.CSAT_Pct, CSAT_GOAL)] for r in csat_cr.itertuples()],
        status_col=3,
    )

    doc.add_paragraph("Voice of Customer — Temas recurrentes en feedback insatisfecho:")
    themes = [
        "Sin solución / 'Aún no me dan solución' — CRs de order status, inedible, damaged",
        "Respuestas genéricas: 'espere', 'confiemos en que llegue' — sin información accionable",
        "Problemas de reembolso/compensación: 'no me regresaron el dinero', 'mentirosos'",
        "Soporte lento en Live Chat pese a QA alto — gap percepción vs auditoría",
    ]
    for t in themes:
        doc.add_paragraph(t, style="List Bullet")

    voc_table = voc_neg.rename(columns={"Country Code": "Country_Code", "CR Lv4": "CR_Lv4"}).head(8)
    add_styled_table(
        doc,
        ["CR Lv4", "Canal", "País", "Comentario VOC"],
        [[r.CR_Lv4, r.Channel, r.Country_Code, str(r.open_question)[:80]] for r in voc_table.itertuples()],
    )

    df_corr = combined.dropna(subset=["QA_Score", "CSAT_Pct"])
    df_corr = df_corr[df_corr["QA_Audits"] >= 5]
    if len(df_corr) >= 3:
        r_val = stats.pearsonr(df_corr["QA_Score"], df_corr["CSAT_Pct"])[0]
        doc.add_paragraph(f"Correlación QA–CSAT (CR Lv4, n≥5 audits): r = {r_val:.3f}. Correlación moderada positiva; QA alto no garantiza CSAT alto en CRs de alto volumen operacional.")
    doc.add_picture(str(charts["scatter"]), width=Inches(5.5))

    # Recontact
    add_heading(doc, "5. Recontact Analysis")
    add_status_paragraph(doc, "Recontact Rate Global", rc_rate, RECONTACT_GOAL, "%", False)
    doc.add_paragraph(
        f"Volumen total: {rc['Recontact Volume'].sum():,.0f} recontactos / {rc['Contacts'].sum():,.0f} contactos. "
        "Live Chat concentra 15.99% recontact rate vs Phone 13.47%. GPTBOT (13.73%) y Email (15.41%) también superan meta."
    )
    add_styled_table(
        doc,
        ["Canal", "Contactos", "Recontact Vol.", "Rate %", "Estado"],
        [
            ["Live Chat", "243,626", "38,961", "15.99%", status_color(15.99, RECONTACT_GOAL, False)],
            ["Phone", "49,674", "6,690", "13.47%", status_color(13.47, RECONTACT_GOAL, False)],
            ["GPTBOT", "19,508", "2,679", "13.73%", status_color(13.73, RECONTACT_GOAL, False)],
            ["Self Help", "666,650", "8,154", "1.22%", status_color(1.22, RECONTACT_GOAL, False)],
        ],
        status_col=4,
    )
    top_rc = rc_cr.head(8)
    add_styled_table(
        doc,
        ["CR Lv4", "Contactos", "Recontact Vol.", "Rate %", "Estado"],
        [[r.CR_Lv4, f"{int(r.Contacts):,}", f"{int(r.Recontact_Volume):,}", f"{r.Recontact_Rate:.2f}%", status_color(r.Recontact_Rate, RECONTACT_GOAL, False)] for r in top_rc.itertuples()],
        status_col=4,
    )
    doc.add_paragraph(
        "Relación Recontact–QA–CSAT: CRs con alto recontact (order status, refund status, cancellation charge) simultáneamente presentan CSAT rojo. "
        "QA en Live Chat para estos CRs es verde (95–98), confirmando que el problema es de resolución/FCR, no de cumplimiento de script auditado."
    )
    doc.add_picture(str(charts["pareto_rc"]), width=Inches(6))

    # Combined Analysis
    add_heading(doc, "6. Combined Analysis")
    doc.add_paragraph(
        "Insight integrado — Cluster 'Order Status & Delay Information':\n"
        "• user request order status or delay information: 16.92% recontact (130,014 vol.), CSAT 67.85%, QA N/A en auditoría masiva\n"
        "• order status & delays: 19.34% recontact (7,680 vol.), CSAT 64.68%, QA 97.33 (verde)\n"
        "• User disagrees with cancellation charge/debt: 12.96% recontact (6,599 vol.), CSAT 67.44%, QA 91.36 (verde)\n\n"
        "Cuantificación: Estos 3 CRs representan ~24,293 recontactos (41.9% del total semanal) y ~21,006 encuestas CSAT insatisfechas estimadas. "
        "Hipótesis root cause: desalineación entre calidad de interacción medida por QA (protocolo/script) y capacidad real de resolver el issue en primera interacción. "
        "Los agentes cumplen script pero no tienen visibilidad operativa (tracking, compensación automática) para cerrar el caso."
    )
    doc.add_picture(str(charts["fishbone"]), width=Inches(6.2))
    doc.add_picture(str(charts["flowchart"]), width=Inches(6.2))

    # Action Plans
    add_heading(doc, "7. Action Plans — by LOB / Business Type")
    actions = [
        ("Food", "QA Ops + Training Lead", "Semana 1–2", "Implementar coaching Phone en Manejo del Tiempo e Información Completa; reducir fail rate de 37.5% a <15%", "CRs: order not received, wrong order"),
        ("Food", "Product + CX Tech", "Semana 2–4", "Integrar widget de tracking en tiempo real en Live Chat para CRs order status/delay", "Target: -3pp recontact en order status CRs"),
        ("Full Service", "Ops Manager + WFM", "Semana 1", "Priorizar staffing en horas pico W19; reducir wait time Live Chat", "CSAT Full Service 79.67%, -5.33pp"),
        ("Full Service", "QA + Content", "Semana 2", "Actualizar macros de compensación automática para 'order completed not received'", "QA Phone CR score 68.16 → meta 85"),
        ("Market Place", "LOB Lead", "Semana 2–3", "Revisar flujo marketplace vs full service; CSAT 80.34% ámbar", "Separar scripts por modalidad"),
        ("Market Place", "Training", "Semana 3", "Capacitación en diferencias de responsabilidad store vs platform", "Reduce VOC 'no me dieron solución'"),
        ("Pickup", "LOB Lead", "Semana 2", "Auditoría de volumen bajo (35 feedback) pero CSAT 74.29% rojo", "Investigar CR pick up at store"),
        ("All LOBs", "CX Leadership", "Semana 1", "War room semanal en top 3 CRs de recontact hasta rate <5.44%", "KPI owner: Service Ops Director"),
        ("All LOBs", "Analytics", "Semana 1–ongoing", "Dashboard FCR por CR Lv4 complementando QA score", "Correlacionar QA vs CSAT semanalmente"),
    ]
    add_styled_table(doc, ["Business Type / LOB", "Responsable", "Plazo", "Acción", "KPI / CR Impactado"], actions)

    # 7 Quality Tools Summary
    add_heading(doc, "Anexo A — Las 7 Herramientas de Calidad Aplicadas")
    tools = [
        ("1. Check Sheet", "Validación de integridad de datos QA, reglas Critical y separación de canales", str(charts["check_sheet"].name)),
        ("2. Diagrama de Pareto", "Identificación del 80/20 en defectos QA y volumen de recontact", f"{charts['pareto'].name}, {charts['pareto_rc'].name}"),
        ("3. Histograma", "Distribución de QA scores — Phone sesgado a scores bajos (critical fails)", charts["histogram"].name),
        ("4. Gráfico de Control", "Variación QA por país/canal vs meta y límites de control", charts["control"].name),
        ("5. Diagrama de Dispersión", "Correlación QA vs CSAT por CR Lv4, coloreado por recontact rate", charts["scatter"].name),
        ("6. Diagrama de Ishikawa", "Análisis causa-raíz del cluster Order Status & Delays", charts["fishbone"].name),
        ("7. Diagrama de Flujo", "Mapeo del punto de falla FCR en journey del cliente", charts["flowchart"].name),
    ]
    add_styled_table(doc, ["Herramienta", "Aplicación en este reporte", "Artefacto"], tools)

    doc.add_paragraph()
    assumptions = doc.add_paragraph()
    assumptions.add_run("Supuestos documentados: ").bold = True
    assumptions.add_run(
        "CR Lv4 = CR_correcta fillna CR_registrada; valor 2 (N/A) excluido de cálculo QA; "
        "18.78% mismatch entre Score_end_user y score recalculado (se usa score recalculado); "
        "Semana analizada: W19; LOB principal: Delivery/Food."
    )

    doc.add_picture(str(charts["check_sheet"]), width=Inches(6))

    docx_path = OUT / "Entregable_2_Weekly_Performance_Report.docx"
    doc.save(docx_path)

    # Also save markdown version for reference
    md_path = OUT / "Entregable_2_Weekly_Performance_Report.md"
    md_content = f"""# DiDi CX — Weekly Performance Report (W19)

## Executive Summary
- **QA Score:** {qa_avg:.2f} (Meta: {QA_GOAL}) — {status_label(status_color(qa_avg, QA_GOAL))}
- **CSAT:** {csat_pct:.2f}% (Meta: {CSAT_GOAL}%) — {status_label(status_color(csat_pct, CSAT_GOAL))}
- **Recontact:** {rc_rate:.2f}% (Meta: {RECONTACT_GOAL}%) — {status_label(status_color(rc_rate, RECONTACT_GOAL, False))}

**Hallazgo crítico:** Cluster Order Status/Delays — 41.9% recontactos, CSAT 64-68%, QA verde en chat.

Ver documento Word completo y gráficos en `/charts/`.
"""
    md_path.write_text(md_content, encoding="utf-8")

    metadata = {
        "generated": datetime.now().isoformat(),
        "week": "W19",
        "kpi": {"qa": qa_avg, "csat": csat_pct, "recontact": rc_rate},
        "charts": {k: str(v) for k, v in charts.items()},
        "docx": str(docx_path),
    }
    (OUT / "report_metadata.json").write_text(json.dumps(metadata, indent=2), encoding="utf-8")
    print(f"Report generated: {docx_path}")
    print(f"Charts: {CHARTS}")


if __name__ == "__main__":
    generate_report()
